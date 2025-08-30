import requests
import json
import os
import sys
from PIL import Image
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

# Tor proxy settings
PROXIES = {
    'http': 'socks5h://127.0.0.1:9050',
    'https': 'socks5h://127.0.0.1:9050'
}

LOG_FILE = "ghost_report.json"

# ==============================
# HTTPS-ONLY BROWSING
# ==============================
def tor_request(url):
    if not url.lower().startswith("https://"):
        return {"status": "blocked", "reason": "Non-HTTPS request blocked"}

    try:
        print(f"[+] Requesting {url} via Tor...")
        r = requests.get(url, proxies=PROXIES, timeout=20)
        return {"status": "ok", "url": url, "code": r.status_code, "preview": r.text[:200]}
    except Exception as e:
        return {"status": "error", "reason": str(e)}

# ==============================
# METADATA CLEANER
# ==============================
def clean_image(infile, outfile):
    try:
        img = Image.open(infile)
        data = list(img.getdata())
        new_img = Image.new(img.mode, img.size)
        new_img.putdata(data)
        new_img.save(outfile)
        return f"[+] Cleaned image saved: {outfile}"
    except Exception as e:
        return f"[!] Error cleaning image: {e}"

def clean_pdf(infile, outfile):
    try:
        reader = PdfReader(infile)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        with open(outfile, "wb") as f:
            writer.write(f)
        return f"[+] Cleaned PDF saved: {outfile}"
    except Exception as e:
        return f"[!] Error cleaning PDF: {e}"

def clean_docx(infile, outfile):
    try:
        doc = Document(infile)
        new_doc = Document()
        for para in doc.paragraphs:
            new_doc.add_paragraph(para.text)
        new_doc.save(outfile)
        return f"[+] Cleaned DOCX saved: {outfile}"
    except Exception as e:
        return f"[!] Error cleaning DOCX: {e}"

# ==============================
# REPORT LOGGER
# ==============================
def save_report(entry):
    try:
        if os.path.exists(LOG_FILE):
            with open(LOG_FILE, "r") as f:
                data = json.load(f)
        else:
            data = []

        data.append(entry)

        with open(LOG_FILE, "w") as f:
            json.dump(data, f, indent=4)
        print(f"[+] Report saved to {LOG_FILE}")
    except Exception as e:
        print(f"[!] Could not save report: {e}")

# ==============================
# MAIN MENU
# ==============================
def main():
    print("\n🕵️ Ghost Protocol: All-in-One Privacy Guard\n")
    print("1. Secure HTTPS request via Tor")
    print("2. Clean file metadata (JPG, PDF, DOCX)")
    print("3. Exit")

    choice = input("Choose: ").strip()

    if choice == "1":
        url = input("Enter HTTPS URL: ").strip()
        result = tor_request(url)
        print(result)
        save_report(result)

    elif choice == "2":
        infile = input("Enter file path: ").strip()
        outfile = input("Enter output file path: ").strip()
        if infile.lower().endswith(".jpg") or infile.lower().endswith(".jpeg") or infile.lower().endswith(".png"):
            print(clean_image(infile, outfile))
        elif infile.lower().endswith(".pdf"):
            print(clean_pdf(infile, outfile))
        elif infile.lower().endswith(".docx"):
            print(clean_docx(infile, outfile))
        else:
            print("[!] Unsupported file type.")

    else:
        print("Exiting...")

if __name__ == "__main__":
    while True:
        print("\n🕵️ Ghost Protocol: All-in-One Privacy Guard\n")
        print("1. Secure HTTPS request via Tor")
        print("2. Clean file metadata (JPG, PDF, DOCX)")
        print("3. Exit")

        choice = input("Choose: ").strip()

        if choice == "1":
            url = input("Enter HTTPS URL: ").strip()
            result = tor_request(url)
            print(result)
            save_report(result)

        elif choice == "2":
            infile = input("Enter file path: ").strip()
            outfile = input("Enter output file path: ").strip()
            if infile.lower().endswith((".jpg", ".jpeg", ".png")):
                print(clean_image(infile, outfile))
            elif infile.lower().endswith(".pdf"):
                print(clean_pdf(infile, outfile))
            elif infile.lower().endswith(".docx"):
                print(clean_docx(infile, outfile))
            else:
                print("[!] Unsupported file type.")

        elif choice == "3":
            print("Exiting...")
            break
        else:
            print("[!] Invalid choice. Try again.")

